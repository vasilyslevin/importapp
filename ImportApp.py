import io
import uuid
import re
import html
import logging
import os  # Add this for environment variables
from typing import List, Dict, Any, Optional

from flask import (
    Flask,
    request,
    jsonify,
    send_file,
    render_template_string,
    abort,
)
from docx import Document
from werkzeug.utils import secure_filename
import google.generativeai as genai  # Add this for Gemini

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB limit

# Configure Gemini API
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

STATE: Dict[str, Dict[str, Any]] = {}
PLACEHOLDER_PATTERN = re.compile(r"\[([^\]]+)\]")


def iter_all_paragraphs(document: Document):
    """Iterate over all paragraphs in the document, including those in tables."""
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def collect_placeholders(document: Document) -> List[str]:
    """Collect unique placeholders from the document."""
    placeholders = []
    seen = set()
    try:
        for paragraph in iter_all_paragraphs(document):
            text = paragraph.text or ""
            for match in PLACEHOLDER_PATTERN.finditer(text):
                name = match.group(1).strip()
                if name and name not in seen:
                    seen.add(name)
                    placeholders.append(name)
    except Exception as e:
        logger.error(f"Error collecting placeholders: {e}")
    return placeholders


def set_paragraph_text(paragraph, text: str):
    """Set the text of a paragraph, preserving the first run if possible."""
    try:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)
    except Exception as e:
        logger.error(f"Error setting paragraph text: {e}")


def apply_values(document: Document, values: Dict[str, str]) -> Document:
    """Apply placeholder values to the document."""
    def substitute(match: re.Match) -> str:
        key = match.group(1).strip()
        return values.get(key, match.group(0))

    try:
        for paragraph in iter_all_paragraphs(document):
            original = paragraph.text or ""
            updated = PLACEHOLDER_PATTERN.sub(substitute, original)
            if updated != original:
                set_paragraph_text(paragraph, updated)
    except Exception as e:
        logger.error(f"Error applying values: {e}")
    return document


def document_to_html(document: Document, values: Optional[Dict[str, str]] = None) -> str:
    """Convert the document to HTML for preview, highlighting unfilled placeholders."""
    parts: List[str] = []
    try:
        for paragraph in document.paragraphs:
            content = paragraph.text or ""
            if values:
                # Highlight unfilled placeholders in red
                content = PLACEHOLDER_PATTERN.sub(
                    lambda m: f'<span style="color: red;">{html.escape(m.group(0))}</span>'
                    if m.group(1).strip() not in values else html.escape(values.get(m.group(1).strip(), m.group(0))),
                    content
                )
            else:
                content = html.escape(content)
            if content.strip():
                parts.append(f"<p>{content}</p>")
        
        for table in document.tables:
            parts.append('<table class="doc-table">')
            for row in table.rows:
                parts.append("<tr>")
                for cell in row.cells:
                    cell_content = []
                    for p in cell.paragraphs:
                        text = p.text or ""
                        if values:
                            text = PLACEHOLDER_PATTERN.sub(
                                lambda m: f'<span style="color: red;">{html.escape(m.group(0))}</span>'
                                if m.group(1).strip() not in values else html.escape(values.get(m.group(1).strip(), m.group(0))),
                                text
                            )
                        else:
                            text = html.escape(text)
                        if text.strip():
                            cell_content.append(text)
                    parts.append(f"<td>{'<br/>'.join(cell_content)}</td>")
                parts.append("</tr>")
            parts.append("</table>")
    except Exception as e:
        logger.error(f"Error converting to HTML: {e}")
        return "<p><em>Error generating preview.</em></p>"
    
    return "".join(parts) or "<p><em>Document is empty.</em></p>"


def get_full_doc_text(document: Document) -> str:
    """Extract full text from the document for AI processing."""
    text_parts = []
    try:
        for paragraph in iter_all_paragraphs(document):
            text = paragraph.text or ""
            if text.strip():
                text_parts.append(text)
    except Exception as e:
        logger.error(f"Error extracting document text: {e}")
    return "\n\n".join(text_parts)


def get_ai_edit(doc_text: str, user_request: str) -> str:
    """Call Gemini API to edit the document based on user request."""
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')  # Use the latest model; adjust if needed
        prompt = f"Document content:\n{doc_text}\n\nUser request: {user_request}\n\nProvide the full edited document text, preserving structure with double newlines for paragraphs."
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        logger.error(f"Error calling Gemini API: {e}")
        return doc_text  # Fallback to original text on error


def apply_ai_edits(document: Document, edited_text: str) -> Document:
    """Apply AI-edited text to the document (flattens to paragraphs; tables are cleared)."""
    try:
        # Clear existing paragraphs
        for para in list(document.paragraphs):
            para.clear()
        # Clear tables (simplified; may lose formatting)
        document._body.clear_content()  # Clears all content
        # Add edited paragraphs
        paragraphs = edited_text.split('\n\n')
        for p_text in paragraphs:
            if p_text.strip():
                document.add_paragraph(p_text.strip())
    except Exception as e:
        logger.error(f"Error applying AI edits: {e}")
    return document


@app.route("/", methods=["GET"])
def index():
    return render_template_string(
        """
<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Legal Template Assistant</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<style>
:root { color-scheme: light dark; font-family: system-ui, sans-serif; }
body { margin: 0; padding: 2rem; background: #f5f5f5; }
main { max-width: 1100px; margin: 0 auto; background: #ffffff; padding: 2rem; border-radius: 12px; box-shadow: 0 10px 30px rgba(0,0,0,0.05); }
h1 { margin-top: 0; }
section { margin-bottom: 2rem; }
label { font-weight: 600; display: block; margin-bottom: 0.5rem; }
input[type="file"] { display: block; margin-bottom: 1rem; }
button { padding: 0.6rem 1.1rem; border: none; border-radius: 6px; background: #2563eb; color: #fff; cursor: pointer; font-weight: 600; }
button:disabled { background: #bcc0c6; cursor: not-allowed; }
.chat-box { border: 1px solid #d0d7de; border-radius: 8px; padding: 1rem; background: #f9f9f9; height: 300px; overflow-y: auto; margin-bottom: 1rem; }
.message { margin-bottom: 1rem; }
.message.assistant { color: #2563eb; }
.message.user { text-align: right; color: #111827; }
#placeholdersList li { padding: 0.4rem 0.6rem; border-radius: 6px; margin-bottom: 0.25rem; background: #eef2ff; font-size: 0.95rem; }
#placeholdersList li.filled { background: #dcfce7; color: #166534; text-decoration: line-through; }
.preview { border: 1px solid #d0d7de; border-radius: 8px; padding: 1rem; background: #fff; max-height: 400px; overflow-y: auto; }
.doc-table { width: 100%; border-collapse: collapse; margin: 1rem 0; }
.doc-table td { border: 1px solid #d0d7de; padding: 0.5rem; }
.status { font-size: 0.9rem; color: #555; margin-top: 0.5rem; }
.hidden { display: none; }
.progress { margin-bottom: 1rem; }
.progress-bar { width: 100%; height: 20px; background: #e5e7eb; border-radius: 10px; overflow: hidden; }
.progress-fill { height: 100%; background: #2563eb; transition: width 0.3s; }
</style>
</head>
<body>
<main>
    <h1>Legal Template Assistant</h1>
    <section>
        <form id="uploadForm">
            <label for="fileInput">Upload a .docx document</label>
            <input id="fileInput" name="file" type="file" accept=".docx" required />
            <button type="submit">Upload</button>
            <p class="status" id="uploadStatus"></p>
        </form>
    </section>

    <section id="placeholdersSection" class="hidden">
        <h2>Template Placeholders</h2>
        <div class="progress" id="progressContainer">
            <div class="progress-bar">
                <div class="progress-fill" id="progressFill" style="width: 0%;"></div>
            </div>
            <p id="progressText">0 of 0 placeholders filled</p>
        </div>
        <ul id="placeholdersList"></ul>
    </section>

    <section id="chatSection" class="hidden">
        <h2>Conversation</h2>
        <div class="chat-box" id="chatMessages"></div>
        <form id="chatForm">
            <input id="chatInput" type="text" placeholder="Type your answer..." autocomplete="off" required />
            <button type="submit">Send</button>
            <button type="button" id="fillAllBtn">Fill All at Once</button>
            <label style="margin-left: 1rem;"><input type="checkbox" id="editMode"> Enable Edit Mode</label>  <!-- Add this -->
        </form>
    </section>

    <section id="previewSection" class="hidden">
        <h2>Preview</h2>
        <div class="preview" id="previewContainer"></div>
        <div style="margin-top:1rem;">
            <button id="refreshPreviewBtn">Refresh Preview</button>
            <a id="downloadLink" class="hidden"><button type="button">Download Completed Document</button></a>
        </div>
    </section>
</main>

<script>
const uploadForm = document.getElementById("uploadForm");
const uploadStatus = document.getElementById("uploadStatus");
const placeholdersSection = document.getElementById("placeholdersSection");
const placeholdersList = document.getElementById("placeholdersList");
const progressContainer = document.getElementById("progressContainer");
const progressFill = document.getElementById("progressFill");
const progressText = document.getElementById("progressText");
const chatSection = document.getElementById("chatSection");
const chatMessages = document.getElementById("chatMessages");
const chatForm = document.getElementById("chatForm");
const chatInput = document.getElementById("chatInput");
const fillAllBtn = document.getElementById("fillAllBtn");
const previewSection = document.getElementById("previewSection");
const previewContainer = document.getElementById("previewContainer");
const downloadLink = document.getElementById("downloadLink");
const refreshPreviewBtn = document.getElementById("refreshPreviewBtn");

let sessionId = null;
let placeholders = [];
let filled = [];

function appendMessage(role, text) {
    const div = document.createElement("div");
    div.classList.add("message", role);
    div.textContent = text;
    chatMessages.appendChild(div);
    chatMessages.scrollTop = chatMessages.scrollHeight;
}

function setPlaceholders(list, filledList) {
    placeholdersList.innerHTML = "";
    filled = filledList;
    list.forEach(name => {
        const li = document.createElement("li");
        li.textContent = name;
        if (filledList.includes(name)) {
            li.classList.add("filled");
        }
        placeholdersList.appendChild(li);
    });
    updateProgress();
}

function updateProgress() {
    const total = placeholders.length;
    const done = filled.length;
    const percent = total > 0 ? (done / total) * 100 : 0;
    progressFill.style.width = `${percent}%`;
    progressText.textContent = `${done} of ${total} placeholders filled`;
}

async function refreshPreview() {
    if (!sessionId) return;
    const res = await fetch(`/preview?session_id=${sessionId}`);
    if (!res.ok) return;
    const payload = await res.json();
    previewContainer.innerHTML = payload.html;
}

uploadForm.addEventListener("submit", async (event) => {
    event.preventDefault();
    if (!uploadForm.file.files.length) return;
    uploadStatus.textContent = "Uploading...";
    const formData = new FormData(uploadForm);
    const res = await fetch("/upload", { method: "POST", body: formData });
    if (!res.ok) {
        const error = await res.json();
        uploadStatus.textContent = error.error || "Upload failed.";
        return;
    }
    const payload = await res.json();
    sessionId = payload.session_id;
    placeholders = payload.placeholders || [];
    setPlaceholders(placeholders, []);
    placeholdersSection.classList.remove("hidden");
    chatSection.classList.toggle("hidden", placeholders.length === 0);
    previewSection.classList.remove("hidden");
    downloadLink.classList.add("hidden");
    downloadLink.removeAttribute("href");
    uploadStatus.textContent = placeholders.length
        ? `Found ${placeholders.length} placeholder(s).`
        : "No placeholders detected.";
    await refreshPreview();
    if (placeholders.length) {
        startConversation();
    }
});

async function startConversation() {
    const res = await fetch("/chat", {
        method: "POST",
        headers: { "Content-Type": "application/json" },
        body: JSON.stringify({ session_id: sessionId, start: true })
    });
    if (!res.ok) return;
    const payload = await res.json();
    handleChatPayload(payload);
}

chatForm.addEventListener("submit", async (event) => {
    event.preventDefault();
    if (!sessionId) return;
    const message = chatInput.value.trim();
    if (!message) return;
    appendMessage("user", message);
    chatInput.value = "";
    const res = await fetch("/chat", {
        method: "POST",
        headers: { "Content-Type": "application/json" },
        body: JSON.stringify({ session_id: sessionId, message, edit_mode: document.getElementById("editMode").checked })  // Add edit_mode
    });
    if (!res.ok) return;
    const payload = await res.json();
    handleChatPayload(payload);
});

fillAllBtn.addEventListener("click", async () => {
    if (!sessionId) return;
    const values = {};
    placeholders.forEach(p => {
        const value = prompt(`Enter value for ${p}:`);
        if (value) values[p] = value;
    });
    if (Object.keys(values).length === 0) return;
    const res = await fetch("/chat", {
        method: "POST",
        headers: { "Content-Type": "application/json" },
        body: JSON.stringify({ session_id: sessionId, fill_all: values })
    });
    if (!res.ok) return;
    const payload = await res.json();
    handleChatPayload(payload);
});

function handleChatPayload(payload) {
    (payload.messages || []).forEach(msg => appendMessage(msg.role, msg.text));
    const filledList = (payload.state && payload.state.filled) || [];
    setPlaceholders(placeholders, filledList);
    if (payload.done) {
        downloadLink.classList.remove("hidden");
        downloadLink.href = `/download?session_id=${sessionId}`;
    }
    refreshPreview();
}

refreshPreviewBtn.addEventListener("click", refreshPreview);
</script>
</body>
</html>
        """
    )


@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "File is required."}), 400
    file = request.files["file"]
    if not file.filename or not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Only .docx files are supported."}), 400
    filename = secure_filename(file.filename)
    contents = file.read()
    if not contents:
        return jsonify({"error": "Uploaded file is empty."}), 400
    try:
        document = Document(io.BytesIO(contents))
        placeholders = collect_placeholders(document)
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        return jsonify({"error": "Unable to process the .docx file."}), 400
    session_id = str(uuid.uuid4())
    STATE[session_id] = {
        "doc_bytes": contents,
        "filename": filename,
        "placeholders": placeholders,
        "values": {},
        "pending": None,
    }
    return jsonify(
        {
            "session_id": session_id,
            "placeholders": placeholders,
            "message": f"Detected {len(placeholders)} placeholder(s)." if placeholders else "No placeholders detected.",
        }
    )


@app.route("/chat", methods=["POST"])
def chat():
    payload = request.get_json(force=True)
    session_id = payload.get("session_id")
    if not session_id or session_id not in STATE:
        return jsonify({"error": "Invalid session."}), 400
    message = (payload.get("message") or "").strip()
    fill_all = payload.get("fill_all")
    start = payload.get("start", False)
    edit_mode = payload.get("edit_mode", False)  # Add this for AI edits
    state = STATE[session_id]
    placeholders = state["placeholders"]
    values = state["values"]
    pending = state.get("pending")
    response_messages = []

    remaining = [p for p in placeholders if p not in values]

    if fill_all:
        values.update(fill_all)
        response_messages.append(
            {"role": "assistant", "text": f"Filled {len(fill_all)} placeholder(s) at once."}
        )
        if not [p for p in placeholders if p not in values]:
            response_messages.append(
                {"role": "assistant", "text": "All placeholders are now filled. You can preview and download the completed document."}
            )
            return jsonify(
                {
                    "messages": response_messages,
                    "done": True,
                    "state": {"filled": list(values.keys()), "pending": None},
                }
            )
        else:
            next_placeholder = [p for p in placeholders if p not in values][0]
            state["pending"] = next_placeholder
            response_messages.append(
                {"role": "assistant", "text": f'Next, provide a value for "[{next_placeholder}]".'}
            )
            return jsonify(
                {
                    "messages": response_messages,
                    "done": False,
                    "state": {"filled": list(values.keys()), "pending": state["pending"]},
                }
            )

    if start:
        if not remaining:
            response_messages.append(
                {"role": "assistant", "text": "No placeholders require input."}
            )
            return jsonify(
                {
                    "messages": response_messages,
                    "done": True,
                    "state": {"filled": list(values.keys()), "pending": None},
                }
            )
        next_placeholder = remaining[0]
        state["pending"] = next_placeholder
        response_messages.append(
            {
                "role": "assistant",
                "text": f'What value should replace "[{next_placeholder}]"?',
            }
        )
        return jsonify(
            {
                "messages": response_messages,
                "done": False,
                "state": {"filled": list(values.keys()), "pending": state["pending"]},
            }
        )

    if edit_mode and message:  # Handle AI edits
        try:
            doc = Document(io.BytesIO(state["doc_bytes"]))
            doc_text = get_full_doc_text(doc)
            edited_text = get_ai_edit(doc_text, message)
            edited_doc = apply_ai_edits(doc, edited_text)
            buffer = io.BytesIO();
            edited_doc.save(buffer)
            state["doc_bytes"] = buffer.getvalue()  # Update stored document
            response_messages.append(
                {"role": "assistant", "text": f"Applied AI edit: {message}"}
            )
        except Exception as e:
            logger.error(f"Error in AI edit: {e}")
            response_messages.append(
                {"role": "assistant", "text": "Sorry, I couldn't apply that edit. Please try again."}
            )
        return jsonify(
            {
                "messages": response_messages,
                "done": False,
                "state": {"filled": list(values.keys()), "pending": pending},
            }
        )

    if pending and message:
        values[pending] = message
        response_messages.append(
            {
                "role": "assistant",
                "text": f'Recorded "{message}" for "[{pending}]".',
            }
        )
        state["pending"] = None
        remaining = [p for p in placeholders if p not in values]
        if remaining:
            next_placeholder = remaining[0]
            state["pending"] = next_placeholder
            response_messages.append(
                {
                    "role": "assistant",
                    "text": f'Next, provide a value for "[{next_placeholder}]".',
                }
            )
            done = False
        else:
            response_messages.append(
                {
                    "role": "assistant",
                    "text": "All placeholders are now filled. You can preview and download the completed document.",
                }
            )
            done = True
        return jsonify(
            {
                "messages": response_messages,
                "done": done,
                "state": {"filled": list(values.keys()), "pending": state["pending"]},
            }
        )

    response_messages.append(
        {"role": "assistant", "text": "Please provide input or check your session."}
    )
    return jsonify(
        {
            "messages": response_messages,
            "done": False,
            "state": {"filled": list(values.keys()), "pending": pending},
        }
    )


@app.route("/preview", methods=["GET"])
def preview():
    session_id = request.args.get("session_id")
    if not session_id or session_id not in STATE:
        return jsonify({"error": "Invalid session."}), 400
    state = STATE[session_id]
    try:
        doc = Document(io.BytesIO(state["doc_bytes"]))
        html_content = document_to_html(doc, state["values"])
    except Exception as e:
        logger.error(f"Error generating preview: {e}")
        return jsonify({"html": "<p><em>Error generating preview.</em></p>"}), 500
    return jsonify({"html": html_content})


@app.route("/download", methods=["GET"])
def download():
    session_id = request.args.get("session_id")
    if not session_id or session_id not in STATE:
        abort(400)
    state = STATE[session_id]
    try:
        doc = Document(io.BytesIO(state["doc_bytes"]))
        filled_doc = apply_values(doc, state["values"])
        buffer = io.BytesIO()
        filled_doc.save(buffer)
        buffer.seek(0)
    except Exception as e:
        logger.error(f"Error preparing download: {e}")
        abort(500)
    original_name = state["filename"]
    base = original_name.rsplit(".", 1)[0]
    download_name = f"{base}_completed.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.teardown_appcontext
def cleanup_sessions(exception=None):
    """Clean up old sessions to prevent memory leaks."""
    global STATE
    # In a real app, implement proper session management with expiration
    pass


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    app.run(host="0.0.0.0", port=port, debug=False) # Set debug=False for production
